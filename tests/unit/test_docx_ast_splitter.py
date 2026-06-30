"""TDD — _DocxASTSplitter: DOCX binary → Document atoms."""

from __future__ import annotations

import io


def _make_docx_bytes(paragraphs: list[tuple[str, str]]) -> bytes:
    """Build a minimal DOCX in-memory. paragraphs is [(style, text), ...]."""
    from docx import Document

    doc = Document()
    # Remove the default empty paragraph
    for para in list(doc.paragraphs):
        p = para._element
        p.getparent().remove(p)

    for style, text in paragraphs:
        if style.startswith("Heading"):
            level = int(style[-1]) if style[-1].isdigit() else 1
            doc.add_heading(text, level=level)
        else:
            doc.add_paragraph(text, style=style if style != "Normal" else None)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _make_docx_with_table(headers: list[str], rows: list[list[str]]) -> bytes:
    from docx import Document

    doc = Document()
    for para in list(doc.paragraphs):
        p = para._element
        p.getparent().remove(p)

    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
    for ri, row in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = val
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _run_splitter(data: bytes) -> list:
    from haystack.dataclasses import Document as HDoc

    from ragent.pipelines.ingest import _DocxASTSplitter

    splitter = _DocxASTSplitter()
    mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    doc = HDoc(content=None, meta={"mime_type": mime, "document_id": "doc-1", "raw_bytes": data})
    return splitter.run([doc])["documents"]


# ---------------------------------------------------------------------------
# Paragraph atoms
# ---------------------------------------------------------------------------


def test_docx_heading_becomes_atom():
    data = _make_docx_bytes([("Heading 1", "Introduction"), ("Normal", "Some body text.")])
    atoms = _run_splitter(data)
    assert len(atoms) == 2
    texts = [a.content for a in atoms]
    assert any("Introduction" in t for t in texts)
    assert any("Some body text." in t for t in texts)


def test_docx_heading_content_strips_markup():
    """content field should be plain text, not include heading markers."""
    data = _make_docx_bytes([("Heading 1", "My Title")])
    atoms = _run_splitter(data)
    assert len(atoms) == 1
    assert atoms[0].content == "My Title"


def test_docx_heading_raw_content_gets_markdown_marker():
    """raw_content must carry a markdown '#'-prefix per heading level so
    downstream heading detection (chat-attachment AST simplification)
    recognizes DOCX headings the same way it recognizes markdown/HTML ones."""
    data = _make_docx_bytes(
        [("Heading 1", "H1 text"), ("Heading 2", "H2 text"), ("Normal", "Body text")]
    )
    atoms = _run_splitter(data)
    assert len(atoms) == 3
    assert atoms[0].meta["raw_content"] == "# H1 text"
    assert atoms[1].meta["raw_content"] == "## H2 text"
    assert atoms[2].meta["raw_content"] == "Body text"


def test_docx_title_style_becomes_level_one_heading():
    from docx import Document

    doc = Document()
    for para in list(doc.paragraphs):
        p = para._element
        p.getparent().remove(p)
    doc.add_heading("My Document", level=0)  # python-docx "Title" style
    buf = io.BytesIO()
    doc.save(buf)
    atoms = _run_splitter(buf.getvalue())
    assert len(atoms) == 1
    assert atoms[0].meta["raw_content"] == "# My Document"


def test_docx_raw_content_set_on_atom():
    data = _make_docx_bytes([("Normal", "Hello world")])
    atoms = _run_splitter(data)
    assert len(atoms) == 1
    assert "raw_content" in atoms[0].meta
    assert "Hello world" in atoms[0].meta["raw_content"]


def test_docx_meta_passthrough():
    """document_id and mime_type from input doc.meta flow to every atom."""
    data = _make_docx_bytes([("Normal", "Text")])
    atoms = _run_splitter(data)
    assert atoms[0].meta["document_id"] == "doc-1"
    mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    assert atoms[0].meta["mime_type"] == mime


def test_docx_empty_document_yields_no_atoms():
    from docx import Document

    doc = Document()
    for para in list(doc.paragraphs):
        p = para._element
        p.getparent().remove(p)
    buf = io.BytesIO()
    doc.save(buf)
    atoms = _run_splitter(buf.getvalue())
    assert atoms == []


def test_docx_blank_paragraphs_skipped():
    data = _make_docx_bytes([("Normal", ""), ("Normal", "Real content"), ("Normal", "")])
    atoms = _run_splitter(data)
    assert len(atoms) == 1
    assert atoms[0].content == "Real content"


# ---------------------------------------------------------------------------
# Table atoms
# ---------------------------------------------------------------------------


def test_docx_table_becomes_atom():
    data = _make_docx_with_table(["Name", "Value"], [["foo", "1"], ["bar", "2"]])
    atoms = _run_splitter(data)
    assert len(atoms) == 1


def test_docx_table_raw_content_is_markdown():
    data = _make_docx_with_table(["Col A", "Col B"], [["x", "y"]])
    atoms = _run_splitter(data)
    raw = atoms[0].meta["raw_content"]
    assert "Col A" in raw
    assert "Col B" in raw
    assert "|" in raw


def test_docx_table_content_is_plain_text():
    """content should be stripped of markdown pipe characters."""
    data = _make_docx_with_table(["H1", "H2"], [["r1c1", "r1c2"]])
    atoms = _run_splitter(data)
    content = atoms[0].content
    assert "H1" in content
    assert "H2" in content


def test_docx_mixed_paragraphs_and_table():
    from docx import Document

    doc = Document()
    for para in list(doc.paragraphs):
        p = para._element
        p.getparent().remove(p)
    doc.add_heading("Overview", level=1)
    doc.add_paragraph("Intro text.")
    table = doc.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "A"
    table.rows[0].cells[1].text = "B"
    table.rows[1].cells[0].text = "1"
    table.rows[1].cells[1].text = "2"
    buf = io.BytesIO()
    doc.save(buf)
    atoms = _run_splitter(buf.getvalue())
    assert len(atoms) == 3  # heading + paragraph + table
